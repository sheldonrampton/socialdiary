topics = {
    "Music and Entertainment": [
        "music", "entertainment", "genres", "musical instruments", "film", "TV", "awards", "events", "artists", "bands", "production"
    ],
    "Technology": [
        "consumer electronics", "software", "applications", "innovations", "digital media", "artificial intelligence", "machine learning", "computers"
    ],
    "Culture and Society": [
        "literature", "arts", "books", "contemporary art", "museums", "social issues", "cultural history", "military history", "fashion", "lifestyle"
    ],
    "Science and Education": [
        "natural sciences", "biology", "chemistry", "ecology", "health", "wellness", "medical", "academia", "higher education", "universities", "research", "environmental science"
    ],
    "Sports and Recreation": [
        "sports", "outdoor activities", "camping", "hiking", "events", "competitions", "olympics", "tennis", "football", "cycling"
    ],
    "Business and Economics": [
        "industry", "markets", "automotive industry", "stock market", "retail", "companies", "brands", "professional development", "career", "entrepreneurship"
    ],
    "Politics and Governance": [
        "political issues", "election", "united states congress", "public policy", "law", "government agencies", "CDC", "environmental policy"
    ],
    "Lifestyle and Personal Interests": [
        "home", "family", "parenting", "home improvement", "personal care", "beauty", "skin care", "hobbies", "crafts", "photography", "arts and crafts", "DIY"
    ]
}

def categorize_post(text, topics):
    for topic, keywords in topics.items():
        if any(keyword in text.lower() for keyword in keywords):
            return topic
    return "Other"

# Example usage
# post_text = "Here is a post about health and wellness."
# category = categorize_post(post_text, topics)
# print(category)  # Outputs: Health


from docx import Document
from datetime import datetime

def create_word_document(posts_data, topics):
    doc = Document()
    doc.add_heading('Facebook Posts', level=1)

    categorized_posts = {topic: [] for topic in topics.keys()}  # Create lists for each topic
    categorized_posts["Other"] = []  # For posts that don't fit any topic

    # Categorize each post
    for post in posts_data:
        if 'text' in post:
            category = categorize_post(post['text'], topics)
            categorized_posts[category].append(post)

    # Write posts by topic to the document
    for category, posts in categorized_posts.items():
        if posts:  # Only add sections for topics with posts
            doc.add_heading(category, level=2)
            for post in posts:
                date = datetime.fromtimestamp(post['timestamp']).strftime('%Y-%m-%d %H:%M:%S')
                doc.add_heading(f'Post Date: {date}', level=3)
                
                if 'text' in post:
                    doc.add_paragraph(f"Post Text: {post['text']}")

                if 'url' in post and 'title' in post:
                    p = doc.add_paragraph()
                    add_hyperlink(p, post['url'], post['title'])

                doc.add_paragraph("")  # Add a blank line for better readability

    doc.save('Facebook_Posts_by_Topic.docx')
    print("Document created successfully!")
